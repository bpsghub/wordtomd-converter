"""DocxConverter: orchestrates .docx → Markdown conversion."""

from __future__ import annotations

import sys
from pathlib import Path
from typing import List, Optional

import docx
from docx.oxml.ns import qn

from wordtomd.numbering import NumberingMap
from wordtomd.postprocess import clean_output
from wordtomd.relationships import RelationshipMap
from wordtomd.renderers.image import ImageExtractor
from wordtomd.renderers.paragraph import render_paragraph
from wordtomd.renderers.table import render_table

_W_TBL = qn("w:tbl")
_W_P = qn("w:p")
_W_SECTPR = qn("w:sectPr")


class DocxConverter:
    def __init__(
        self,
        input_path: Path,
        output_path: Path,
        image_dir: Optional[str] = None,
        extract_images: bool = True,
        verbose: bool = False,
    ) -> None:
        self.input_path = input_path
        self.output_path = output_path
        self.extract_images = extract_images
        self.verbose = verbose

        # Derive images directory next to the output file
        images_dir_name = image_dir or (output_path.stem + "_images")
        self.images_dir = output_path.parent / images_dir_name

    def _log(self, msg: str) -> None:
        if self.verbose:
            print(f"[wordtomd] {msg}", file=sys.stderr)

    def convert(self) -> None:
        self._log(f"Opening {self.input_path}")
        rel_map = RelationshipMap.from_docx(str(self.input_path))
        num_map = NumberingMap.from_docx(str(self.input_path))
        doc = docx.Document(str(self.input_path))

        image_extractor = ImageExtractor(
            rel_map=rel_map,
            images_dir=self.images_dir,
            enabled=self.extract_images,
        )

        output_lines: List[str] = []
        last_block_type: str = "empty"
        last_num_id: Optional[str] = None
        code_buffer: List[str] = []

        def flush_code_buffer() -> None:
            nonlocal last_block_type
            if code_buffer:
                if output_lines:
                    output_lines.append("")
                output_lines.append("```")
                output_lines.extend(code_buffer)
                output_lines.append("```")
                code_buffer.clear()
                last_block_type = "code"

        def emit_block(block_type: str, lines: List[str], num_id: Optional[str] = None) -> None:
            nonlocal last_block_type, last_num_id
            if not lines:
                return

            # Add blank line separator:
            # - between any non-list block and the next block
            # - when exiting a list
            # - when switching between different list groups (different numId)
            needs_blank = False
            if output_lines:
                if last_block_type not in ("empty", "list"):
                    needs_blank = True
                elif last_block_type == "list" and block_type != "list":
                    needs_blank = True
                elif last_block_type == "list" and block_type == "list" and num_id and num_id != last_num_id:
                    needs_blank = True

            if needs_blank:
                output_lines.append("")

            output_lines.extend(lines)
            last_block_type = block_type
            if block_type == "list":
                last_num_id = num_id

        # Walk body children directly to get both paragraphs and tables in order
        body = doc.element.body
        for child in body:
            tag = child.tag

            if tag == _W_SECTPR:
                continue  # section properties — skip

            if tag == _W_TBL:
                # Find the matching docx Table object
                flush_code_buffer()
                # python-docx stores tables in doc.tables; match by element
                tbl_obj = None
                for t in doc.tables:
                    if t._tbl is child:
                        tbl_obj = t
                        break
                if tbl_obj is not None:
                    tbl_lines = render_table(tbl_obj, rel_map)
                    if tbl_lines:
                        emit_block("table", tbl_lines)
                continue

            if tag == _W_P:
                # Find the matching docx Paragraph object
                para_obj = None
                for p in doc.paragraphs:
                    if p._p is child:
                        para_obj = p
                        break
                if para_obj is None:
                    # Paragraph inside a table cell — skip here (handled by table renderer)
                    continue

                block_type, lines = render_paragraph(para_obj, rel_map, num_map, image_extractor)

                if block_type == "code":
                    # Buffer consecutive code paragraphs
                    code_buffer.extend(lines)
                else:
                    flush_code_buffer()
                    # Resolve the numId for list group boundary detection
                    para_num_id: Optional[str] = None
                    if block_type == "list":
                        from wordtomd.renderers.list_item import _get_num_pr
                        para_num_id, _ = _get_num_pr(para_obj)
                    emit_block(block_type, lines, num_id=para_num_id)

        # Flush any remaining code buffer
        flush_code_buffer()

        result = clean_output(output_lines)

        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.output_path.write_text(result, encoding="utf-8")
        self._log(f"Written to {self.output_path}")
